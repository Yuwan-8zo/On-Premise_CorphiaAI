"""
文件服務模組

處理文件上傳、解析、分塊和向量化
"""

import logging
import os
import uuid
from pathlib import Path
from typing import Optional, BinaryIO

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.config import settings
from app.models.document import Document, DocumentStatus
from app.models.document_chunk import DocumentChunk
from app.services.rag_service import get_rag_service

logger = logging.getLogger(__name__)


class DocumentService:
    """文件處理服務"""
    
    # 支援的檔案類型
    SUPPORTED_TYPES = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }
    
    # 分塊設定
    CHUNK_SIZE = 500  # 字元數
    CHUNK_OVERLAP = 50  # 重疊字元數
    
    def __init__(self, db: AsyncSession):
        self.db = db
        self.upload_dir = Path(settings.upload_directory)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    async def upload_document(
        self,
        file: BinaryIO,
        filename: str,
        tenant_id: str,
        user_id: str,
    ) -> Document:
        """
        上傳並儲存文件

        支援版本化：同一租戶內重新上傳相同原始檔名的文件時，
        舊版會被標記 superseded_by → 新版 ID，新版 version 自動遞增。
        
        Args:
            file: 檔案物件
            filename: 原始檔名
            tenant_id: 租戶 ID
            user_id: 上傳者 ID
            
        Returns:
            Document: 文件記錄
        """
        # 檢查檔案類型
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_TYPES:
            raise ValueError(f"不支援的檔案類型: {ext}")
        
        # 生成唯一檔名
        unique_filename = f"{uuid.uuid4()}{ext}"
        file_path = self.upload_dir / tenant_id / unique_filename
        file_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 儲存檔案
        content = file.read()
        file_size = len(content)
        
        with open(file_path, "wb") as f:
            f.write(content)
        
        # B3: 查詢同名舊版本（同一租戶、同原始檔名、尚未被取代的最新版）
        prev_result = await self.db.execute(
            select(Document)
            .where(
                Document.tenant_id == tenant_id,
                Document.original_filename == filename,
                Document.superseded_by.is_(None),  # 只找「當前版本」
            )
            .order_by(Document.version.desc())
            .limit(1)
        )
        prev_doc = prev_result.scalar_one_or_none()
        new_version = (prev_doc.version + 1) if prev_doc else 1
        
        # 建立資料庫記錄
        document = Document(
            tenant_id=tenant_id,
            uploaded_by=user_id,
            filename=unique_filename,
            original_filename=filename,
            file_type=ext[1:],  # 移除點號
            file_size=file_size,
            file_path=str(file_path),
            status=DocumentStatus.PENDING.value,
            version=new_version,
        )
        
        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)
        
        # B3: 標記舊版被取代
        if prev_doc:
            prev_doc.superseded_by = document.id
            await self.db.commit()
            logger.info(
                f"文件版本更新: {filename} v{prev_doc.version} → v{new_version}"
            )
        
        logger.info(f"已上傳文件: {filename} → {unique_filename} (v{new_version})")
        
        return document
    
    async def process_document(self, document_id: str) -> bool:
        """
        處理文件：解析、分塊、向量化
        
        Args:
            document_id: 文件 ID
            
        Returns:
            bool: 是否成功
        """
        # 取得文件記錄
        result = await self.db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if document is None:
            logger.error(f"文件不存在: {document_id}")
            return False
        
        # 更新狀態
        document.status = DocumentStatus.PROCESSING.value
        await self.db.commit()
        
        try:
            # 解析文件內容
            content = await self._parse_document(document)
            
            if not content:
                raise ValueError("無法解析文件內容")
            
            # 分塊
            chunks = self._chunk_text(content)
            
            rag_service = get_rag_service()
            
            # 儲存分塊到資料庫
            for i, chunk_content in enumerate(chunks):
                # 取得向量
                embedding = await rag_service.get_embedding(chunk_content)
                
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk_content,
                    embedding=embedding,
                    chunk_metadata={
                        "filename": document.original_filename,
                        "tenant_id": document.tenant_id,
                    }
                )
                self.db.add(chunk)
            
            # 更新文件狀態
            document.status = DocumentStatus.COMPLETED.value
            document.chunk_count = len(chunks)
            # NOTE: 使用 naive UTC，與系統其他地方保持一致
            from app.core.time_utils import utc_now_naive
            document.processed_at = utc_now_naive()
            
            await self.db.commit()
            
            logger.info(f"文件處理完成: {document.original_filename}，共 {len(chunks)} 個分塊")
            return True
            
        except Exception as e:
            logger.error(f"文件處理失敗: {e}")
            document.status = DocumentStatus.FAILED.value
            document.error_message = str(e)
            await self.db.commit()
            return False
    
    async def _parse_document(self, document: Document) -> str:
        """解析文件內容"""
        file_path = Path(document.file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"檔案不存在: {file_path}")
        
        ext = document.file_type.lower()
        
        if ext in ["txt", "md"]:
            return await self._parse_text_file(file_path)
        elif ext == "pdf":
            return await self._parse_pdf(file_path)
        elif ext in ["docx", "doc"]:
            return await self._parse_word(file_path)
        elif ext in ["xlsx", "xls"]:
            return await self._parse_excel(file_path)
        else:
            raise ValueError(f"不支援的檔案類型: {ext}")
    
    async def _parse_text_file(self, file_path: Path) -> str:
        """解析純文字檔案"""
        import chardet
        import asyncio
        
        def _process():
            with open(file_path, "rb") as f:
                raw_data = f.read()
            # 偵測編碼
            detected = chardet.detect(raw_data)
            encoding = detected.get("encoding", "utf-8")
            return raw_data.decode(encoding, errors="ignore")
            
        return await asyncio.to_thread(_process)
    
    async def _parse_pdf(self, file_path: Path) -> str:
        """解析 PDF 檔案"""
        import asyncio
        
        def _process():
            try:
                from PyPDF2 import PdfReader
                
                reader = PdfReader(str(file_path))
                text_parts = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return "\n\n".join(text_parts)
                
            except ImportError:
                logger.warning("PyPDF2 未安裝，無法解析 PDF")
                raise
                
        return await asyncio.to_thread(_process)
    
    async def _parse_word(self, file_path: Path) -> str:
        """解析 Word 檔案"""
        import asyncio
        
        def _process():
            try:
                from docx import Document as DocxDocument
                
                doc = DocxDocument(str(file_path))
                text_parts = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                return "\n\n".join(text_parts)
                
            except ImportError:
                logger.warning("python-docx 未安裝，無法解析 Word")
                raise
                
        return await asyncio.to_thread(_process)
    
    async def _parse_excel(self, file_path: Path) -> str:
        """解析 Excel 檔案"""
        import asyncio
        
        def _process():
            try:
                from openpyxl import load_workbook
                
                wb = load_workbook(str(file_path), data_only=True)
                text_parts = []
                
                for sheet in wb.worksheets:
                    sheet_text = f"### {sheet.title}\n"
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join(str(cell) if cell else "" for cell in row)
                        if row_text.strip():
                            sheet_text += row_text + "\n"
                    text_parts.append(sheet_text)
                
                return "\n\n".join(text_parts)
                
            except ImportError:
                logger.warning("openpyxl 未安裝，無法解析 Excel")
                raise
                
        return await asyncio.to_thread(_process)
    
    def _chunk_text(self, text: str) -> list[str]:
        """
        將文字分塊
        
        ISSUE-03 修正：使用字元數切割，支援中文文本（中文無空白分隔詞）。
        
        Args:
            text: 原始文字
            
        Returns:
            list[str]: 分塊列表
        """
        chunks = []
        
        # 先按段落分割
        paragraphs = text.split("\n\n")
        
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 如果段落本身超過 chunk size，需要進一步分割
            if len(para) > self.CHUNK_SIZE:
                # 先保存當前累積的 chunk
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                
                # ISSUE-03 修正：使用字元切割，而非 split()（空白斷詞對中文無效）
                # 以 CHUNK_SIZE 個字元為單位切割，加入 CHUNK_OVERLAP 字元的重疊避免語意截斷
                start = 0
                while start < len(para):
                    end = min(start + self.CHUNK_SIZE, len(para))
                    chunk_part = para[start:end]
                    if chunk_part.strip():
                        chunks.append(chunk_part)
                    # 下一段起點往回退 CHUNK_OVERLAP 個字元，保留語意連貫性
                    start = end - self.CHUNK_OVERLAP if end < len(para) else end
            else:
                # 嘗試合併段落
                if len(current_chunk) + len(para) + 2 > self.CHUNK_SIZE:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = para
                else:
                    current_chunk = current_chunk + "\n\n" + para if current_chunk else para
        
        # 加入最後一個 chunk
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    async def delete_document(self, document_id: str) -> bool:
        """
        刪除文件
        
        Args:
            document_id: 文件 ID
            
        Returns:
            bool: 是否成功
        """
        result = await self.db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()
        
        if document is None:
            return False
        
        try:
            # 刪除檔案
            file_path = Path(document.file_path)
            if file_path.exists():
                os.remove(file_path)
            
            # 由於 DocumentChunk 設定了 CASCADE，刪除 Document 時會一併從資料庫刪除所有向量分塊
            
            # 刪除資料庫記錄
            await self.db.delete(document)
            await self.db.commit()
            
            logger.info(f"已刪除文件: {document.original_filename}")
            return True
            
        except Exception as e:
            logger.error(f"刪除文件失敗: {e}")
            return False
